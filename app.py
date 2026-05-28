import os
import io
import json
import uuid
import time
import urllib.request
from functools import wraps
from datetime import datetime

from flask import (
    Flask, render_template, request, redirect, url_for,
    session, flash, send_file, Response, abort
)
from dotenv import load_dotenv
from supabase import create_client, Client
from docx import Document

load_dotenv()

app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "dev-secret-change-me")
app.wsgi_app = __import__('whitenoise').WhiteNoise(app.wsgi_app, root='static/', prefix='static')

SUPABASE_URL = os.environ.get("SUPABASE_URL", "")
SUPABASE_KEY = os.environ.get("SUPABASE_KEY", "")
SUPABASE_SERVICE_KEY = os.environ.get("SUPABASE_SERVICE_KEY", "")
SUPABASE_BUCKET = os.environ.get("SUPABASE_BUCKET", "docgen-files")

supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)
supabase_admin: Client = create_client(SUPABASE_URL, SUPABASE_SERVICE_KEY) if SUPABASE_SERVICE_KEY else supabase

# Alle database-queries via de service_role client zodat RLS-policies
# op alle tabellen ingeschakeld kunnen worden zonder toegangsproblemen.
db = supabase_admin

# ---------------------------------------------------------------------------
# Config helpers
# ---------------------------------------------------------------------------

def get_config() -> dict:
    """Load tenant config from Supabase config table, with env var fallbacks."""
    defaults = {
        "tenant_name": os.environ.get("TENANT_NAME", "Mijn Organisatie"),
        "primary_color": os.environ.get("PRIMARY_COLOR", "#2563EB"),
        "logo_url": "",
    }
    try:
        rows = db.table("config").select("key,value").execute()
        if rows.data:
            for row in rows.data:
                defaults[row["key"]] = row["value"]
    except Exception:
        pass
    return defaults


def set_config(key: str, value: str) -> None:
    db.table("config").upsert({"key": key, "value": value}).execute()


_footer_cache: dict = {"data": None, "ts": 0}
_FOOTER_TTL = 300  # 5 minuten

def get_app_footer() -> dict:
    """Haal footer-config op van het intranet, gecached voor 5 minuten."""
    now = time.time()
    if _footer_cache["data"] is not None and now - _footer_cache["ts"] < _FOOTER_TTL:
        return _footer_cache["data"]

    url = os.environ.get(
        "FOOTER_API_URL",
        "https://intranet.leidersinzicht.nl/api/v1/app-footers/docgen"
    )
    try:
        with urllib.request.urlopen(url, timeout=3) as resp:
            data = json.loads(resp.read().decode())
        if data.get("jaar_dynamisch"):
            data["tekst"] = data.get("tekst", "").replace("{jaar}", str(datetime.utcnow().year))
        _footer_cache["data"] = data
        _footer_cache["ts"]   = now
        return data
    except Exception:
        fallback = {"tekst": "", "url": "", "url_label": "", "afbeelding_url": ""}
        return fallback


def get_ongelezen_inzendingen_count(user_id: str) -> int:
    """Aantal door externe partij verzonden invullingen na de laatste keer dat de gebruiker de inzendingen-pagina bezocht."""
    try:
        gezien_res = db.table("inzendingen_gelezen").select("gezien_op").eq("user_id", user_id).single().execute()
        gezien_op = gezien_res.data["gezien_op"] if gezien_res.data else None
    except Exception:
        gezien_op = None

    try:
        q = db.table("invullingen").select("id,updated_at").eq("extern_status", "verzonden")
        if gezien_op:
            q = q.gt("updated_at", gezien_op)
        inv_res = q.execute()
        return len(inv_res.data or [])
    except Exception:
        return 0


# ---------------------------------------------------------------------------
# Auth helpers
# ---------------------------------------------------------------------------

@app.context_processor
def inject_globals():
    user_id = session.get("user_id")
    return {
        "ongelezen_inzendingen": get_ongelezen_inzendingen_count(user_id) if user_id else 0,
        "app_footer": get_app_footer(),
    }


def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get("user_id"):
            return redirect(url_for("login"))
        return f(*args, **kwargs)
    return decorated


def admin_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get("user_id"):
            return redirect(url_for("login"))
        if session.get("user_role") != "admin":
            flash("Je hebt geen toegang tot dit onderdeel.", "error")
            return redirect(url_for("index"))
        return f(*args, **kwargs)
    return decorated


# ---------------------------------------------------------------------------
# Word generation
# ---------------------------------------------------------------------------

def _merge_placeholder_runs(para, placeholder: str, key: str):
    """
    Als een placeholder gesplitst is over meerdere runs, normaliseer die naar één run.
    De opmaak van de run die de veldnaam bevat (niet {{ of }}) wordt gebruikt.
    Geeft de gemerge run terug, of None als de placeholder niet gesplitst aanwezig is.
    """
    runs = para.runs
    full_text = "".join(r.text for r in runs)
    if placeholder not in full_text:
        return None

    # Zoek de aaneengesloten reeks runs die samen de placeholder bevatten
    start_idx = end_idx = None
    accumulated = ""
    for i, run in enumerate(runs):
        accumulated += run.text
        if placeholder in accumulated:
            end_idx = i
            # Ga terug om start te vinden
            acc2 = ""
            for j in range(i, -1, -1):
                acc2 = runs[j].text + acc2
                if placeholder in acc2:
                    start_idx = j
                    break
            break

    if start_idx is None:
        return None

    # Kies als carrier de run die de veldnaam bevat — niet {{ of }}
    # Dat voorkomt dat de opmaak van {{ (die soms bold is als decoratie) erft.
    carrier_idx = start_idx
    for i in range(start_idx, end_idx + 1):
        if key in runs[i].text:
            carrier_idx = i
            break

    # Schrijf de volledige tekst van het bereik samen in de carrier-run,
    # leeg de overige runs in het bereik
    merged_text = "".join(r.text for r in runs[start_idx:end_idx + 1])
    runs[carrier_idx].text = merged_text
    for i in range(start_idx, end_idx + 1):
        if i != carrier_idx:
            runs[i].text = ""

    return runs[carrier_idx]


def _replace_in_paragraph(para, values: dict):
    """Replace placeholders in a paragraph, even when split across multiple runs."""
    for key, val in values.items():
        placeholder = f"{{{{{key}}}}}"
        if placeholder not in para.text:
            continue
        val_str = str(val) if val is not None else ""
        # Fast path: placeholder zit volledig in één run
        replaced = False
        for run in para.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, val_str)
                replaced = True
                break
        if replaced:
            continue
        # Slow path: normaliseer gesplitste runs naar één carrier-run, dan vervang
        carrier = _merge_placeholder_runs(para, placeholder, key)
        if carrier is not None:
            carrier.text = carrier.text.replace(placeholder, val_str)


SYSTEM_VARIABLES = [
    ("_datum_vandaag",      "Datum van aanmaken document (bijv. 12 mei 2026)"),
    ("_datum_iso",          "Datum in ISO-formaat (bijv. 2026-05-12)"),
    ("_jaar",               "Huidig jaar (bijv. 2026)"),
    ("_jaar-1",             "Vorig jaar (bijv. 2025)"),
    ("_jaar+1",             "Volgend jaar (bijv. 2027)"),
    ("_maand",              "Huidige maand voluit (bijv. mei)"),
    ("_tenant_naam",        "Naam van de organisatie uit de instellingen"),
    ("_ingevuld_door",      "Label van de toegangscode waarmee ingelogd is"),
    ("_dossier_naam",       "Naam van het dossier (alleen bij dossier-download)"),
    ("_dossier_omschrijving", "Omschrijving van het dossier (alleen bij dossier-download)"),
    ("_bijlagen_lijst",     "Kommalijst van alle sjabloonnamen in het dossier"),
    ("_bijlagen_aantal",    "Aantal sjablonen in het dossier"),
    ("_bijlage_volgnummer", "Volgnummer van dit sjabloon binnen het dossier (bijv. 2)"),
]

MAANDEN = ["januari","februari","maart","april","mei","juni",
           "juli","augustus","september","oktober","november","december"]

def build_system_values(cfg: dict, dossier: dict = None, templates_in_dossier: list = None, positie: int = None, user_label: str = None) -> dict:
    now = datetime.now()
    sv = {
        "_datum_vandaag": f"{now.day} {MAANDEN[now.month - 1]} {now.year}",
        "_datum_iso":     now.strftime("%Y-%m-%d"),
        "_jaar":          str(now.year),
        "_jaar-1":        str(now.year - 1),
        "_jaar+1":        str(now.year + 1),
        "_maand":         MAANDEN[now.month - 1],
        "_tenant_naam":   cfg.get("tenant_name", ""),
        "_ingevuld_door": user_label or "",
        "_dossier_naam":          (dossier or {}).get("naam", ""),
        "_dossier_omschrijving":  (dossier or {}).get("omschrijving", ""),
        "_bijlagen_lijst":    ", ".join(t.get("name", "") for t in (templates_in_dossier or [])),
        "_bijlagen_aantal":   str(len(templates_in_dossier)) if templates_in_dossier is not None else "",
        "_bijlage_volgnummer": str(positie) if positie is not None else "",
    }
    # Financieringsvormen: elke actieve vorm levert een _fin_<slug> systeemvariabele op.
    # Waarde is "ja" als de vorm actief is op het dossier, anders "" (zodat {{#if}} werkt).
    actieve_vormen = set()
    if dossier:
        raw = dossier.get("financieringsvorm") or ""
        actieve_vormen = {v.strip() for v in raw.split(",") if v.strip()}
    for naam in _get_financieringsvormen():
        sv[_fin_slug(naam)] = "ja" if naam in actieve_vormen else ""
    return sv


DATE_FORMATS = {
    "DD-MM-YYYY":   "%d-%m-%Y",
    "DD-MM-YY":     "%d-%m-%y",
    "D MMM YYYY":   None,   # speciaal: Nederlandse afkorting
    "D MMMM YYYY":  None,   # speciaal: Nederlandse voluit
    "YYYY-MM-DD":   "%Y-%m-%d",
}

MAANDEN_KORT = ["jan","feb","mrt","apr","mei","jun","jul","aug","sep","okt","nov","dec"]

def _format_date(iso_str: str, fmt_key: str) -> str:
    from datetime import datetime as dt
    try:
        d = dt.strptime(iso_str, "%Y-%m-%d")
    except ValueError:
        return iso_str
    if fmt_key == "D MMM YYYY":
        return f"{d.day} {MAANDEN_KORT[d.month - 1]} {d.year}"
    if fmt_key == "D MMMM YYYY":
        return f"{d.day} {MAANDEN[d.month - 1]} {d.year}"
    strfmt = DATE_FORMATS.get(fmt_key, "%d-%m-%Y")
    return d.strftime(strfmt)


def format_field_values(fields: list, values: dict) -> dict:
    """Zet datum-velden om naar gewenst formaat en laat overige waarden intact."""
    result = dict(values)
    for field in fields:
        key = field.get("name")
        if field.get("type") == "date" and key in result and result[key]:
            fmt_key = field.get("date_format", "DD-MM-YYYY")
            result[key] = _format_date(result[key], fmt_key)
    return result


import re as _re
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy
import lxml.etree as _etree

_IF_RE    = _re.compile(r'\{\{#if\s+(\w+)\}\}')
_IFNOT_RE = _re.compile(r'\{\{#ifnot\s+(\w+)\}\}')
_ENDIF_RE = _re.compile(r'\{\{/if(?:not)?\}\}')


def _para_full_text(para) -> str:
    return "".join(r.text for r in para.runs)


def _para_set_text(para, text: str):
    """Overschrijf alle runs: eerste run krijgt volledige tekst, rest leeg.
    Verwijdert ook w:br elementen uit alle runs én als directe kind van w:p."""
    # Verwijder w:br direct onder w:p
    for br in para._element.findall(qn('w:br')):
        para._element.remove(br)
    # Verwijder w:br binnen w:r elementen
    for r_elem in para._element.findall(qn('w:r')):
        for br in r_elem.findall(qn('w:br')):
            r_elem.remove(br)
    if para.runs:
        para.runs[0].text = text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(text)


def _remove_para_spacing(p_elem):
    """Zet w:spacing after/before op 0 in de pPr van een alinea-element."""
    pPr = p_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p_elem.insert(0, pPr)
    # Verwijder bestaand spacing-element en vervang met schone versie
    for old in pPr.findall(qn('w:spacing')):
        pPr.remove(old)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:after'), '0')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:line'), '240')
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)
    # contextualSpacing voorkomt extra ruimte tussen alinea's van dezelfde stijl (o.a. lijsten)
    for old in pPr.findall(qn('w:contextualSpacing')):
        pPr.remove(old)
    ctx = OxmlElement('w:contextualSpacing')
    ctx.set(qn('w:val'), '0')
    pPr.append(ctx)


def _insert_paragraph_after(ref_para, text: str, source_para, remove_spacing: bool = False):
    """Voeg een nieuwe alinea in direct na ref_para, met alinea-opmaak van source_para."""
    from docx.text.paragraph import Paragraph as DocxParagraph

    # Kopieer de alinea-structuur (inclusief pPr/opmaak), maar zonder runs en w:br
    new_p = deepcopy(source_para._element)
    for r in new_p.findall(qn('w:r')):
        new_p.remove(r)
    for br in new_p.findall(qn('w:br')):
        new_p.remove(br)

    if remove_spacing:
        _remove_para_spacing(new_p)

    # Bouw een nieuwe run met de tekst
    new_r = OxmlElement('w:r')
    # Kopieer run-opmaak van eerste run als die er is
    if source_para.runs:
        rpr = source_para.runs[0]._element.find(qn('w:rPr'))
        if rpr is not None:
            new_r.append(deepcopy(rpr))
    new_t = OxmlElement('w:t')
    new_t.text = text
    new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    new_r.append(new_t)
    new_p.append(new_r)

    ref_para._element.addnext(new_p)
    return DocxParagraph(new_p, source_para._element.getparent())


def _strip_tag_from_para(p_elem, pattern):
    """Verwijder alle tekst die matcht met pattern uit de runs van een alinea-element.
    Werkt ook als de tag gesplitst is over meerdere w:t elementen."""
    t_elems = list(p_elem.iter(qn('w:t')))
    full = "".join(t.text or "" for t in t_elems)
    if not pattern.search(full):
        return
    new_full = pattern.sub('', full)
    # Schrijf gecombineerde tekst naar het eerste t-element, leeg de rest
    if t_elems:
        t_elems[0].text = new_full
        t_elems[0].set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        for t in t_elems[1:]:
            t.text = ""


def _process_conditionals(paragraphs_parent, values: dict):
    """
    Verwerk {{#if veld}}, {{#ifnot veld}}, {{/if}} blokken.

    Ondersteunt twee patronen:
    1. Tag + inhoud + sluitingstag op dezelfde alinea:
       {{#if veld}}tekst{{/if}}  →  alinea bewaren of verwijderen als geheel
    2. Openingstag op alinea A, sluitingstag op (of na) alinea B:
       {{#if veld}}          ← alinea A: altijd verwijderen
       ...tussenliggende alineas...
       tekst{{/if}}          ← alinea B: sluitingstag strippen, alinea bewaren/verwijderen
    """
    body = paragraphs_parent
    paras = [child for child in body if child.tag == qn('w:p')]

    to_delete = set()
    i = 0
    while i < len(paras):
        p = paras[i]
        text = "".join(t.text or "" for t in p.iter(qn('w:t')))

        m_if    = _IF_RE.search(text)
        m_ifnot = _IFNOT_RE.search(text)

        if m_if or m_ifnot:
            field  = (m_if or m_ifnot).group(1)
            is_if  = bool(m_if)
            filled = bool(values.get(field, "").strip())
            keep   = filled if is_if else not filled

            # Controleer of openings- én sluitingstag op dezelfde alinea staan
            open_tag_re = m_if.re if m_if else m_ifnot.re
            if _ENDIF_RE.search(text):
                # Patroon 1: alles op één alinea — strip beide tags of verwijder alinea
                if keep:
                    _strip_tag_from_para(p, open_tag_re)
                    _strip_tag_from_para(p, _ENDIF_RE)
                else:
                    to_delete.add(id(p))
                i += 1
            else:
                # Patroon 2: openingstag-alinea altijd verwijderen
                to_delete.add(id(p))
                i += 1

                depth = 1
                while i < len(paras) and depth > 0:
                    inner = paras[i]
                    inner_text = "".join(t.text or "" for t in inner.iter(qn('w:t')))

                    has_open  = bool(_IF_RE.search(inner_text) or _IFNOT_RE.search(inner_text))
                    has_close = bool(_ENDIF_RE.search(inner_text))

                    if has_open:
                        depth += 1
                    if has_close:
                        depth -= 1
                        if depth == 0:
                            # Sluitingstag: strip de tag, bewaar of verwijder de alinea
                            if keep:
                                _strip_tag_from_para(inner, _ENDIF_RE)
                            else:
                                to_delete.add(id(inner))
                            i += 1
                            break
                    if not keep and not has_close:
                        to_delete.add(id(inner))
                    i += 1
        else:
            i += 1

    for p in paras:
        if id(p) in to_delete:
            p.getparent().remove(p)


def _process_textarea(para, values: dict):
    """
    Als een alinea een textarea-placeholder bevat met meerdere regels (\n),
    vervang de alinea door meerdere alinea's (één per regel).
    Anders: gewone placeholder-vervanging.
    """
    # Herstel de volledige tekst inclusief gesplitste runs
    full_text = "".join(r.text or "" for r in para.runs)

    for key, val in values.items():
        placeholder = f"{{{{{key}}}}}"
        if placeholder not in full_text:
            continue
        val_str = str(val) if val is not None else ""
        # Normaliseer Windows CRLF naar LF
        val_str = val_str.replace("\r\n", "\n").replace("\r", "\n")
        if "\n" not in val_str:
            continue
        # Multiline: trim afsluitende lege regels, splits op newlines
        lines = val_str.rstrip("\n").split("\n")
        # Verwijder spacing van de originele alinea (w:br wordt door _para_set_text verwijderd)
        _remove_para_spacing(para._element)
        first_line_text = full_text.replace(placeholder, lines[0])
        _para_set_text(para, first_line_text)
        # Volgende regels als nieuwe alinea's, allemaal zonder spacing
        prev = para
        for line in lines[1:]:
            line_text = full_text.replace(placeholder, line)
            prev = _insert_paragraph_after(prev, line_text, para, remove_spacing=True)
        return  # klaar, sla gewone vervanging over

    # Geen multiline — gewone vervanging
    _replace_in_paragraph(para, values)


def fill_template(docx_bytes: bytes, values: dict) -> bytes:
    doc = Document(io.BytesIO(docx_bytes))

    # Normaliseer Windows CRLF in alle waarden (browsers sturen \r\n vanuit textarea)
    values = {k: v.replace("\r\n", "\n").replace("\r", "\n") if isinstance(v, str) else v
              for k, v in values.items()}

    # Stap 1: conditionele blokken verwerken (verwijdert alinea's)
    _process_conditionals(doc.element.body, values)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _process_conditionals(cell._tc, values)

    # Stap 2: placeholders vervangen (inclusief multiline textarea)
    for para in doc.paragraphs:
        _process_textarea(para, values)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _process_textarea(para, values)

    # Stap 3: placeholders in kop- en voetteksten per sectie
    for section in doc.sections:
        for hdr_ftr in (
            section.header, section.footer,
            section.even_page_header, section.even_page_footer,
            section.first_page_header, section.first_page_footer,
        ):
            if hdr_ftr is None:
                continue
            for para in hdr_ftr.paragraphs:
                _process_textarea(para, values)
            for table in hdr_ftr.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            _process_textarea(para, values)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Auth routes
# ---------------------------------------------------------------------------

@app.route("/login", methods=["GET", "POST"])
def login():
    cfg = get_config()
    if request.method == "POST":
        email = request.form.get("email", "").strip()
        password = request.form.get("password", "")
        if not email or not password:
            flash("Voer je e-mailadres en wachtwoord in.", "error")
            return render_template("login.html", cfg=cfg)

        try:
            auth_res = supabase.auth.sign_in_with_password({"email": email, "password": password})
        except Exception as e:
            app.logger.error(f"Supabase login fout: {e}")
            flash("Ongeldige inloggegevens.", "error")
            return render_template("login.html", cfg=cfg)

        user = auth_res.user if auth_res else None
        app.logger.info(f"Login poging email={email!r} user={user}")
        if user:
            metadata = user.user_metadata or {}
            session["user_id"] = str(user.id)
            session["user_email"] = user.email
            session["user_role"] = metadata.get("role", "user")
            session["user_name"] = metadata.get("name", "")
            return redirect(url_for("index"))
        else:
            flash("Ongeldige inloggegevens.", "error")

    return render_template("login.html", cfg=cfg)


@app.route("/logout")
def logout():
    try:
        supabase.auth.sign_out()
    except Exception:
        pass
    session.clear()
    return redirect(url_for("login"))


# ---------------------------------------------------------------------------
# Internal routes
# ---------------------------------------------------------------------------

@app.route("/")
@login_required
def index():
    return redirect(url_for("dossiers_overzicht"))


@app.route("/template/<template_id>")
@login_required
def template_detail(template_id):
    cfg = get_config()
    try:
        tmpl = db.table("templates").select("*").eq("id", template_id).single().execute()
    except Exception:
        abort(404)

    if not tmpl.data:
        abort(404)

    template = tmpl.data
    fields = template.get("fields") or []
    if isinstance(fields, str):
        fields = json.loads(fields)

    try:
        tokens_res = db.table("tokens").select("*").eq("template_id", template_id).order("created_at", desc=True).execute()
        tokens = tokens_res.data or []
    except Exception:
        tokens = []

    return render_template(
        "template_detail.html",
        cfg=cfg,
        template=template,
        fields=fields,
        tokens=tokens,
    )


@app.route("/template/<template_id>/download", methods=["POST"])
@login_required
def template_download(template_id):
    try:
        tmpl = db.table("templates").select("*").eq("id", template_id).single().execute()
    except Exception:
        abort(404)

    if not tmpl.data:
        abort(404)

    template = tmpl.data
    fields = template.get("fields") or []
    if isinstance(fields, str):
        fields = json.loads(fields)

    values = {}
    for field in fields:
        key = field["name"]
        values[key] = request.form.get(key, "")

    values = format_field_values(fields, values)
    cfg = get_config()
    system_vals = build_system_values(cfg, user_label=session.get("user_name", session.get("user_email", "")))
    values = {**system_vals, **values}

    try:
        docx_bytes = supabase.storage.from_(SUPABASE_BUCKET).download(template["docx_path"])
    except Exception as e:
        flash(f"Fout bij ophalen sjabloonbestand: {e}", "error")
        return redirect(url_for("template_detail", template_id=template_id))

    try:
        result = fill_template(docx_bytes, values)
    except Exception as e:
        flash(f"Fout bij genereren document: {e}", "error")
        return redirect(url_for("template_detail", template_id=template_id))

    filename = f"{template['name'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return send_file(
        io.BytesIO(result),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename,
    )


@app.route("/template/<template_id>/tokens/create", methods=["POST"])
@login_required
def token_create(template_id):
    description = request.form.get("description", "").strip() or "Extern formulier"
    try:
        db.table("tokens").insert({
            "template_id": template_id,
            "description": description,
            "status": "pending",
        }).execute()
        flash("Token aangemaakt.", "success")
    except Exception as e:
        flash(f"Fout bij aanmaken token: {e}", "error")
    return redirect(url_for("template_detail", template_id=template_id))


@app.route("/token/<token_id>/unseal", methods=["POST"])
@login_required
def token_unseal(token_id):
    try:
        # Fetch token to get template_id for redirect
        tok = db.table("tokens").select("template_id").eq("id", token_id).single().execute()
        db.table("tokens").update({"status": "pending"}).eq("id", token_id).execute()
        flash("Token heropend — formulier kan opnieuw worden ingevuld.", "success")
        if tok.data:
            return redirect(url_for("template_detail", template_id=tok.data["template_id"]))
    except Exception as e:
        flash(f"Fout bij heropenen token: {e}", "error")
    return redirect(url_for("index"))


@app.route("/token/<token_id>/download")
@login_required
def token_download(token_id):
    try:
        tok = db.table("tokens").select("*").eq("id", token_id).single().execute()
    except Exception:
        abort(404)

    if not tok.data:
        abort(404)

    token = tok.data
    if token["status"] != "sealed":
        flash("Dit token heeft nog geen ingediende invulling.", "error")
        return redirect(url_for("template_detail", template_id=token["template_id"]))

    try:
        sub = db.table("submissions").select("*").eq("token_id", token_id).order("submitted_at", desc=True).limit(1).single().execute()
    except Exception:
        flash("Geen inzending gevonden voor dit token.", "error")
        return redirect(url_for("template_detail", template_id=token["template_id"]))

    if not sub.data:
        flash("Geen inzending gevonden voor dit token.", "error")
        return redirect(url_for("template_detail", template_id=token["template_id"]))

    values = sub.data.get("values") or {}
    if isinstance(values, str):
        values = json.loads(values)

    try:
        tmpl = db.table("templates").select("*").eq("id", token["template_id"]).single().execute()
    except Exception:
        abort(404)

    template = tmpl.data
    t_fields = template.get("fields") or []
    if isinstance(t_fields, str):
        t_fields = json.loads(t_fields)
    values = format_field_values(t_fields, values)
    cfg = get_config()
    system_vals = build_system_values(cfg, user_label=session.get("user_name", session.get("user_email", "")))
    values = {**system_vals, **values}

    try:
        docx_bytes = supabase.storage.from_(SUPABASE_BUCKET).download(template["docx_path"])
    except Exception as e:
        flash(f"Fout bij ophalen sjabloonbestand: {e}", "error")
        return redirect(url_for("template_detail", template_id=token["template_id"]))

    result = fill_template(docx_bytes, values)
    desc = token.get("description", "document").replace(" ", "_")
    filename = f"{desc}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return send_file(
        io.BytesIO(result),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename,
    )


# ---------------------------------------------------------------------------
# External routes (no login)
# ---------------------------------------------------------------------------

@app.route("/fill/<token_id>", methods=["GET", "POST"])
def fill_external(token_id):
    cfg = get_config()
    try:
        tok = db.table("tokens").select("*").eq("id", token_id).single().execute()
    except Exception:
        abort(404)

    if not tok.data:
        abort(404)

    token = tok.data

    if token["status"] == "sealed":
        return render_template("fill_thanks.html", cfg=cfg, already_sealed=True)

    try:
        tmpl = db.table("templates").select("*").eq("id", token["template_id"]).single().execute()
    except Exception:
        abort(404)

    if not tmpl.data:
        abort(404)

    template = tmpl.data
    fields = template.get("fields") or []
    if isinstance(fields, str):
        fields = json.loads(fields)

    if request.method == "POST":
        values = {}
        errors = []
        for field in fields:
            key = field["name"]
            val = request.form.get(key, "").strip()
            if field.get("required") and not val:
                errors.append(f'"{field.get("label", key)}" is verplicht.')
            values[key] = val

        if errors:
            for err in errors:
                flash(err, "error")
            return render_template(
                "fill_external.html",
                cfg=cfg,
                template=template,
                fields=fields,
                token=token,
                prefill=values,
            )

        try:
            db.table("submissions").insert({
                "token_id": token_id,
                "values": values,
            }).execute()
            db.table("tokens").update({"status": "sealed"}).eq("id", token_id).execute()
        except Exception as e:
            flash(f"Fout bij opslaan: {e}", "error")
            return render_template(
                "fill_external.html",
                cfg=cfg,
                template=template,
                fields=fields,
                token=token,
                prefill=values,
            )

        return redirect(url_for("fill_thanks", token_id=token_id))

    return render_template(
        "fill_external.html",
        cfg=cfg,
        template=template,
        fields=fields,
        token=token,
        prefill={},
    )


@app.route("/fill/<token_id>/thanks")
def fill_thanks(token_id):
    cfg = get_config()
    try:
        tok = db.table("tokens").select("status").eq("id", token_id).single().execute()
        already_sealed = tok.data and tok.data.get("status") == "sealed"
    except Exception:
        already_sealed = True
    return render_template("fill_thanks.html", cfg=cfg, already_sealed=already_sealed)


# ---------------------------------------------------------------------------
# Admin routes
# ---------------------------------------------------------------------------

@app.route("/sjablonen")
@login_required
def sjablonen():
    cfg = get_config()
    try:
        templates_res = db.table("templates").select("id,name,created_at").order("created_at", desc=True).execute()
        templates = templates_res.data or []
    except Exception:
        templates = []
    fin_vormen = [{"naam": n, "slug": _fin_slug(n)} for n in _get_financieringsvormen()]
    return render_template("sjablonen.html", cfg=cfg, templates=templates, system_variables=SYSTEM_VARIABLES, fin_vormen=fin_vormen)


@app.route("/admin")
@admin_required
def admin():
    cfg = get_config()
    try:
        templates_res = db.table("templates").select("id,name,created_at").order("created_at", desc=True).execute()
        templates = templates_res.data or []
    except Exception:
        templates = []

    try:
        tokens_res = db.table("tokens").select("id,status").execute()
        token_stats = {
            "total": len(tokens_res.data or []),
            "sealed": sum(1 for t in (tokens_res.data or []) if t["status"] == "sealed"),
            "pending": sum(1 for t in (tokens_res.data or []) if t["status"] == "pending"),
        }
    except Exception:
        token_stats = {"total": 0, "sealed": 0, "pending": 0}

    fin_vormen = [{"naam": n, "slug": _fin_slug(n)} for n in _get_financieringsvormen()]
    return render_template(
        "admin.html",
        cfg=cfg,
        templates=templates,
        token_stats=token_stats,
        system_variables=SYSTEM_VARIABLES,
        fin_vormen=fin_vormen,
    )


@app.route("/admin/template/scan", methods=["POST"])
@login_required
def admin_template_scan():
    """Scan een geüpload .docx bestand en geef gevonden placeholders terug als JSON."""
    docx_file = request.files.get("docx_file")
    if not docx_file or not docx_file.filename.lower().endswith(".docx"):
        return {"error": "Geen geldig .docx bestand."}, 400

    try:
        doc = Document(io.BytesIO(docx_file.read()))
    except Exception as e:
        return {"error": f"Kon bestand niet lezen: {e}"}, 400

    import re
    found = set()
    system_keys = {key for key, _ in SYSTEM_VARIABLES}

    def scan_text(text):
        for match in re.findall(r"\{\{([^}]+)\}\}", text):
            key = match.strip()
            # Sla systeemvariabelen en conditionele tags over
            if not key:
                continue
            if key in system_keys:
                continue
            if key.startswith('#if ') or key.startswith('#ifnot ') or key.startswith('/if'):
                continue
            found.add(key)

    for para in doc.paragraphs:
        scan_text(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    scan_text(para.text)

    return {"placeholders": sorted(found)}


@app.route("/admin/template/new", methods=["GET", "POST"])
@login_required
def admin_template_new():
    cfg = get_config()
    if request.method == "POST":
        name = request.form.get("name", "").strip()
        description = request.form.get("description", "").strip()
        fields_json = request.form.get("fields_json", "[]").strip()
        docx_file = request.files.get("docx_file")

        if not name:
            flash("Naam is verplicht.", "error")
            return render_template("admin_template_edit.html", cfg=cfg, template=None, mode="new")

        if not docx_file or not docx_file.filename:
            flash("Upload een .docx sjabloonbestand.", "error")
            return render_template("admin_template_edit.html", cfg=cfg, template=None, mode="new")

        if not docx_file.filename.lower().endswith(".docx"):
            flash("Alleen .docx bestanden zijn toegestaan.", "error")
            return render_template("admin_template_edit.html", cfg=cfg, template=None, mode="new")

        try:
            fields = json.loads(fields_json)
        except json.JSONDecodeError:
            fields = []

        file_bytes = docx_file.read()
        storage_path = f"templates/{uuid.uuid4()}.docx"

        try:
            supabase.storage.from_(SUPABASE_BUCKET).upload(
                storage_path,
                file_bytes,
                file_options={"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
            )
        except Exception as e:
            flash(f"Fout bij uploaden bestand: {e}", "error")
            return render_template("admin_template_edit.html", cfg=cfg, template=None, mode="new")

        try:
            db.table("templates").insert({
                "name": name,
                "description": description,
                "docx_path": storage_path,
                "fields": fields,
            }).execute()
            flash("Sjabloon aangemaakt.", "success")
            return redirect(url_for("admin"))
        except Exception as e:
            flash(f"Fout bij opslaan sjabloon: {e}", "error")

    return render_template("admin_template_edit.html", cfg=cfg, template=None, mode="new")


@app.route("/admin/template/<template_id>/edit", methods=["GET", "POST"])
@login_required
def admin_template_edit(template_id):
    cfg = get_config()
    try:
        tmpl = db.table("templates").select("*").eq("id", template_id).single().execute()
    except Exception:
        abort(404)

    if not tmpl.data:
        abort(404)

    template = tmpl.data
    fields = template.get("fields") or []
    if isinstance(fields, str):
        fields = json.loads(fields)

    if request.method == "POST":
        name = request.form.get("name", "").strip()
        description = request.form.get("description", "").strip()
        fields_json = request.form.get("fields_json", "[]").strip()
        docx_file = request.files.get("docx_file")

        if not name:
            flash("Naam is verplicht.", "error")
            return render_template("admin_template_edit.html", cfg=cfg, template=template, fields=fields, mode="edit")

        try:
            new_fields = json.loads(fields_json)
        except json.JSONDecodeError:
            new_fields = fields

        update_data = {
            "name": name,
            "description": description,
            "fields": new_fields,
            "updated_at": datetime.utcnow().isoformat(),
        }

        if docx_file and docx_file.filename and docx_file.filename.lower().endswith(".docx"):
            file_bytes = docx_file.read()
            storage_path = f"templates/{uuid.uuid4()}.docx"
            try:
                supabase.storage.from_(SUPABASE_BUCKET).upload(
                    storage_path,
                    file_bytes,
                    file_options={"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
                )
                # Delete old file
                try:
                    supabase.storage.from_(SUPABASE_BUCKET).remove([template["docx_path"]])
                except Exception:
                    pass
                update_data["docx_path"] = storage_path
            except Exception as e:
                flash(f"Fout bij uploaden nieuw bestand: {e}", "error")
                return render_template("admin_template_edit.html", cfg=cfg, template=template, fields=fields, mode="edit")

        try:
            db.table("templates").update(update_data).eq("id", template_id).execute()
            flash("Sjabloon bijgewerkt.", "success")
            return redirect(url_for("admin"))
        except Exception as e:
            flash(f"Fout bij opslaan: {e}", "error")

    return render_template("admin_template_edit.html", cfg=cfg, template=template, fields=fields, mode="edit")


@app.route("/admin/template/<template_id>/delete", methods=["POST"])
@login_required
def admin_template_delete(template_id):
    try:
        tmpl = db.table("templates").select("docx_path").eq("id", template_id).single().execute()
        if tmpl.data:
            try:
                supabase.storage.from_(SUPABASE_BUCKET).remove([tmpl.data["docx_path"]])
            except Exception:
                pass
        db.table("templates").delete().eq("id", template_id).execute()
        flash("Sjabloon verwijderd.", "success")
    except Exception as e:
        flash(f"Fout bij verwijderen: {e}", "error")
    return redirect(url_for("admin"))


@app.route("/admin/config", methods=["GET", "POST"])
@admin_required
def admin_config():
    cfg = get_config()

    if request.method == "POST":
        action = request.form.get("action")

        if action == "save_config":
            tenant_name = request.form.get("tenant_name", "").strip()
            primary_color = request.form.get("primary_color", "").strip()
            logo_url = request.form.get("logo_url", "").strip()

            if tenant_name:
                set_config("tenant_name", tenant_name)
            if primary_color:
                set_config("primary_color", primary_color)
            set_config("logo_url", logo_url)
            flash("Configuratie opgeslagen.", "success")
            return redirect(url_for("admin_config"))

        if action == "add_user":
            new_name = request.form.get("new_name", "").strip()
            new_email = request.form.get("new_email", "").strip()
            new_password = request.form.get("new_password", "").strip()
            new_role = request.form.get("new_role", "user")
            if new_role not in ("admin", "user"):
                new_role = "user"

            if not new_email or not new_password:
                flash("E-mailadres en wachtwoord zijn verplicht.", "error")
            else:
                try:
                    supabase_admin.auth.admin.create_user({
                        "email": new_email,
                        "password": new_password,
                        "user_metadata": {"name": new_name, "role": new_role},
                        "email_confirm": True,
                    })
                    flash("Gebruiker aangemaakt.", "success")
                except Exception as e:
                    flash(f"Fout: {e}", "error")
            return redirect(url_for("admin_config"))

        if action == "delete_user":
            user_id = request.form.get("user_id")
            try:
                supabase_admin.auth.admin.delete_user(user_id)
                flash("Gebruiker verwijderd.", "success")
            except Exception as e:
                flash(f"Fout: {e}", "error")
            return redirect(url_for("admin_config"))

        if action == "add_financieringsvorm":
            naam = request.form.get("fin_naam", "").strip()
            if naam:
                try:
                    db.table("financieringsvormen").insert({"naam": naam}).execute()
                    flash("Financieringsvorm toegevoegd.", "success")
                except Exception as e:
                    flash(f"Fout: {e}", "error")
            else:
                flash("Naam is verplicht.", "error")
            return redirect(url_for("admin_config"))

        if action == "rename_financieringsvorm":
            fin_id = request.form.get("fin_id")
            naam = request.form.get("fin_naam", "").strip()
            if fin_id and naam:
                try:
                    db.table("financieringsvormen").update({"naam": naam}).eq("id", fin_id).execute()
                    flash("Financieringsvorm bijgewerkt.", "success")
                except Exception as e:
                    flash(f"Fout: {e}", "error")
            return redirect(url_for("admin_config"))

        if action == "delete_financieringsvorm":
            fin_id = request.form.get("fin_id")
            if fin_id:
                try:
                    db.table("financieringsvormen").delete().eq("id", fin_id).execute()
                    flash("Financieringsvorm verwijderd.", "success")
                except Exception as e:
                    flash(f"Fout: {e}", "error")
            return redirect(url_for("admin_config"))

        if action == "add_dossier_type":
            naam = request.form.get("dt_naam", "").strip()
            beschrijving = request.form.get("dt_beschrijving", "").strip()
            if naam:
                try:
                    db.table("dossier_types").insert({"naam": naam, "beschrijving": beschrijving or None}).execute()
                    flash("Dossiertype toegevoegd.", "success")
                except Exception as e:
                    flash(f"Fout: {e}", "error")
            else:
                flash("Naam is verplicht.", "error")
            return redirect(url_for("admin_config"))

        if action == "update_dossier_type":
            dt_id = request.form.get("dt_id")
            naam = request.form.get("dt_naam", "").strip()
            beschrijving = request.form.get("dt_beschrijving", "").strip()
            if dt_id and naam:
                try:
                    db.table("dossier_types").update({"naam": naam, "beschrijving": beschrijving or None}).eq("id", dt_id).execute()
                    flash("Dossiertype bijgewerkt.", "success")
                except Exception as e:
                    flash(f"Fout: {e}", "error")
            return redirect(url_for("admin_config"))

        if action == "delete_dossier_type":
            dt_id = request.form.get("dt_id")
            if dt_id:
                try:
                    db.table("dossier_types").delete().eq("id", dt_id).execute()
                    flash("Dossiertype verwijderd.", "success")
                except Exception as e:
                    flash(f"Fout: {e}", "error")
            return redirect(url_for("admin_config"))

    try:
        users_res = supabase_admin.auth.admin.list_users()
        raw = users_res if isinstance(users_res, list) else []
        users = [{
            "id": str(u.id),
            "email": u.email or "",
            "created_at": str(u.created_at or ""),
            "user_metadata": u.user_metadata or {},
        } for u in raw]
    except Exception:
        users = []

    try:
        fin_res = db.table("financieringsvormen").select("*").order("naam").execute()
        financieringsvormen = fin_res.data or []
        for f in financieringsvormen:
            f["slug"] = _fin_slug(f["naam"])
    except Exception:
        financieringsvormen = []

    dossier_types = _get_dossier_types()

    return render_template("admin_config.html", cfg=cfg, users=users, financieringsvormen=financieringsvormen, dossier_types=dossier_types)



# ---------------------------------------------------------------------------
# Dossier routes (ingelogd vereist)
# ---------------------------------------------------------------------------

@app.route("/dossiers")
@login_required
def dossiers_overzicht():
    cfg = get_config()
    try:
        dos_res = db.table("dossiers").select("*").order("created_at", desc=True).execute()
        dossier_list = dos_res.data or []
    except Exception as e:
        flash(f"Fout bij ophalen dossiers: {e}", "error")
        dossier_list = []

    # Enrich with invulling count
    for dos in dossier_list:
        try:
            cnt_res = db.table("invullingen").select("id").eq("dossier_id", dos["id"]).execute()
            dos["invulling_count"] = len(cnt_res.data or [])
        except Exception:
            dos["invulling_count"] = 0

    # Haal weergavevoorkeur op
    user_id = session.get("user_id")
    view_mode = "kaarten"
    try:
        pref_res = db.table("user_preferences").select("preferences").eq("user_id", user_id).execute()
        if pref_res.data:
            view_mode = pref_res.data[0]["preferences"].get("dossiers_view", "kaarten")
    except Exception:
        pass

    return render_template("dossiers.html", cfg=cfg, dossiers=dossier_list, view_mode=view_mode)


def _get_dossier_types() -> list:
    """Geeft lijst van {id, naam, beschrijving} dossier-types terug, of [] als geen geconfigureerd."""
    try:
        res = db.table("dossier_types").select("id,naam,beschrijving").order("naam").execute()
        return res.data or []
    except Exception:
        return []


def _fin_slug(naam: str) -> str:
    """Zet een financieringsvorm-naam om naar een template-placeholder slug, bijv. 'Wlz' → '_fin_wlz'."""
    import re as _re
    slug = naam.lower().strip()
    slug = _re.sub(r'[^a-z0-9]+', '_', slug).strip('_')
    return f"_fin_{slug}"


def _get_financieringsvormen() -> list:
    try:
        res = db.table("financieringsvormen").select("naam").order("naam").execute()
        return [r["naam"] for r in (res.data or [])]
    except Exception:
        return ["Zvw", "Wlz", "Wmo", "Jeugdwet", "Overig"]


@app.route("/dossier/nieuw", methods=["GET", "POST"])
@login_required
def dossier_nieuw():
    cfg = get_config()
    try:
        tmpl_res = db.table("templates").select("id,name,description").order("name").execute()
        templates = tmpl_res.data or []
    except Exception:
        templates = []
    fin_vormen = _get_financieringsvormen()
    dossier_types = _get_dossier_types()

    if request.method == "POST":
        naam = request.form.get("naam", "").strip()
        omschrijving = request.form.get("omschrijving", "").strip()
        template_ids = request.form.getlist("template_ids")
        jaar_raw = request.form.get("jaar", "").strip()
        vormen = request.form.getlist("financieringsvorm")
        financieringsvorm = ", ".join(sorted(v.strip() for v in vormen if v.strip())) or None
        oa_type = request.form.get("oa_type", "").strip() or None

        try:
            jaar = int(jaar_raw) if jaar_raw else None
        except ValueError:
            jaar = None

        if not naam:
            flash("Naam is verplicht.", "error")
            return render_template("dossier_nieuw.html", cfg=cfg, templates=templates, now=datetime.now(), fin_vormen=fin_vormen, dossier_types=dossier_types)

        if not template_ids:
            flash("Selecteer minimaal één sjabloon.", "error")
            return render_template("dossier_nieuw.html", cfg=cfg, templates=templates, now=datetime.now(), fin_vormen=fin_vormen, dossier_types=dossier_types)

        try:
            dos_res = db.table("dossiers").insert({
                "naam": naam,
                "omschrijving": omschrijving or None,
                "jaar": jaar,
                "financieringsvorm": financieringsvorm,
                "oa_type": oa_type,
            }).execute()
            dossier_id = dos_res.data[0]["id"]
        except Exception as e:
            flash(f"Fout bij aanmaken dossier: {e}", "error")
            return render_template("dossier_nieuw.html", cfg=cfg, templates=templates, now=datetime.now(), fin_vormen=fin_vormen, dossier_types=dossier_types)

        for tid in template_ids:
            try:
                db.table("invullingen").insert({
                    "dossier_id": dossier_id,
                    "template_id": tid,
                    "waarden": {},
                    "extern_toegang": "verborgen",
                }).execute()
            except Exception as e:
                flash(f"Fout bij aanmaken invulling: {e}", "error")

        flash("Dossier aangemaakt.", "success")
        return redirect(url_for("dossier_detail", dossier_id=dossier_id))

    return render_template("dossier_nieuw.html", cfg=cfg, templates=templates, now=datetime.now(), fin_vormen=fin_vormen, dossier_types=dossier_types)


@app.route("/dossier/<dossier_id>")
@login_required
def dossier_detail(dossier_id):
    cfg = get_config()
    try:
        dos_res = db.table("dossiers").select("*").eq("id", dossier_id).single().execute()
    except Exception:
        abort(404)

    if not dos_res.data:
        abort(404)

    dossier = dos_res.data

    try:
        inv_res = db.table("invullingen").select("*").eq("dossier_id", dossier_id).execute()
        invullingen_raw = inv_res.data or []
    except Exception:
        invullingen_raw = []

    # Enrich invullingen with template field definitions
    invullingen = []
    for inv in invullingen_raw:
        try:
            tmpl_res = db.table("templates").select("*").eq("id", inv["template_id"]).single().execute()
            tmpl = tmpl_res.data or {}
        except Exception:
            tmpl = {}
        fields = tmpl.get("fields") or []
        if isinstance(fields, str):
            fields = json.loads(fields)
        waarden = inv.get("waarden") or {}
        if isinstance(waarden, str):
            waarden = json.loads(waarden)
        invullingen.append({
            **inv,
            "template": tmpl,
            "fields": fields,
            "waarden": waarden,
            "is_filled": any(str(v).strip() for v in waarden.values()),
        })

    try:
        tok_res = db.table("dossier_tokens").select("*").eq("dossier_id", dossier_id).order("created_at", desc=True).execute()
        tokens = tok_res.data or []
    except Exception:
        tokens = []

    fin_vormen = _get_financieringsvormen()
    dossier_types = _get_dossier_types()

    try:
        all_templates_res = db.table("templates").select("id,name").order("name").execute()
        all_templates = all_templates_res.data or []
    except Exception:
        all_templates = []

    # Verzamel unieke dossier-scope velden uit alle gekoppelde templates (op volgorde van eerste optreden)
    raw_gw = dossier.get("gedeelde_waarden")
    if raw_gw is None:
        gedeelde_waarden = {}
    elif isinstance(raw_gw, str):
        try:
            gedeelde_waarden = json.loads(raw_gw)
        except (json.JSONDecodeError, ValueError):
            gedeelde_waarden = {}
    else:
        gedeelde_waarden = dict(raw_gw)  # zorg dat het een gewone Python dict is
    seen_names = set()
    dossier_velden = []
    for inv in invullingen:
        for field in inv.get("fields", []):
            if field.get("scope") == "dossier" and field["name"] not in seen_names:
                seen_names.add(field["name"])
                dossier_velden.append(field)

    return render_template(
        "dossier_detail.html",
        cfg=cfg,
        dossier=dossier,
        invullingen=invullingen,
        tokens=tokens,
        fin_vormen=fin_vormen,
        all_templates=all_templates,
        dossier_velden=dossier_velden,
        gedeelde_waarden=gedeelde_waarden,
        dossier_types=dossier_types,
    )


@app.route("/dossier/<dossier_id>/invulling/<inv_id>", methods=["POST"])
@login_required
def dossier_invulling_opslaan(dossier_id, inv_id):
    try:
        inv_res = db.table("invullingen").select("*").eq("id", inv_id).eq("dossier_id", dossier_id).single().execute()
    except Exception:
        abort(404)

    if not inv_res.data:
        abort(404)

    inv = inv_res.data
    try:
        tmpl_res = db.table("templates").select("fields").eq("id", inv["template_id"]).single().execute()
    except Exception:
        abort(404)

    fields = tmpl_res.data.get("fields") or []
    if isinstance(fields, str):
        fields = json.loads(fields)

    # Start from existing waarden to preserve extern-filled values
    existing_waarden = inv.get("waarden") or {}
    if isinstance(existing_waarden, str):
        import json as _json
        existing_waarden = _json.loads(existing_waarden)

    waarden = dict(existing_waarden)
    for field in fields:
        if field.get("scope") == "dossier":
            continue  # dossier-scope velden worden opgeslagen op dossier-niveau, niet per invulling
        key = field["name"]
        waarden[key] = request.form.get(key, "")

    try:
        db.table("invullingen").update({
            "waarden": waarden,
            "updated_at": datetime.utcnow().isoformat(),
        }).eq("id", inv_id).execute()
        flash("Invulling opgeslagen.", "success")
    except Exception as e:
        flash(f"Fout bij opslaan: {e}", "error")

    return redirect(url_for("dossier_detail", dossier_id=dossier_id))


@app.route("/dossier/<dossier_id>/gedeelde-waarden", methods=["POST"])
@login_required
def dossier_gedeelde_waarden_opslaan(dossier_id):
    try:
        dos_res = db.table("dossiers").select("gedeelde_waarden").eq("id", dossier_id).single().execute()
    except Exception:
        abort(404)
    if not dos_res.data:
        abort(404)

    existing = dos_res.data.get("gedeelde_waarden") or {}
    if isinstance(existing, str):
        existing = json.loads(existing)

    # Verwerk alle form-velden die met "dv_" beginnen (dossier-veld prefix)
    nieuwe_waarden = dict(existing)
    for key, val in request.form.items():
        if key.startswith("dv_"):
            field_name = key[3:]  # strip "dv_" prefix
            nieuwe_waarden[field_name] = val

    try:
        db.table("dossiers").update({
            "gedeelde_waarden": nieuwe_waarden,
            "updated_at": datetime.utcnow().isoformat(),
        }).eq("id", dossier_id).execute()
        flash("Gedeelde gegevens opgeslagen.", "success")
    except Exception as e:
        flash(f"Fout bij opslaan: {e}", "error")

    return redirect(url_for("dossier_detail", dossier_id=dossier_id))


@app.route("/dossier/<dossier_id>/sjabloon-toevoegen", methods=["POST"])
@login_required
def dossier_sjabloon_toevoegen(dossier_id):
    try:
        dos_res = db.table("dossiers").select("id").eq("id", dossier_id).single().execute()
    except Exception:
        abort(404)
    if not dos_res.data:
        abort(404)

    template_ids = request.form.getlist("template_ids")
    if not template_ids:
        flash("Selecteer minimaal één sjabloon.", "error")
        return redirect(url_for("dossier_detail", dossier_id=dossier_id))

    # Haal bestaande template_ids op om duplicaten te voorkomen
    try:
        existing_res = db.table("invullingen").select("template_id").eq("dossier_id", dossier_id).execute()
        existing_ids = {str(r["template_id"]) for r in (existing_res.data or [])}
    except Exception:
        existing_ids = set()

    added = 0
    for tid in template_ids:
        if str(tid) in existing_ids:
            continue
        try:
            db.table("invullingen").insert({
                "dossier_id": dossier_id,
                "template_id": tid,
                "waarden": {},
                "extern_toegang": "verborgen",
            }).execute()
            added += 1
        except Exception as e:
            flash(f"Fout bij toevoegen sjabloon: {e}", "error")

    if added:
        flash(f"{added} sjabloon{'en' if added != 1 else ''} toegevoegd.", "success")
    else:
        flash("Geselecteerde sjablonen waren al gekoppeld aan dit dossier.", "warning")

    return redirect(url_for("dossier_detail", dossier_id=dossier_id))


@app.route("/dossier/<dossier_id>/invulling/<inv_id>/download")
@login_required
def dossier_invulling_download(dossier_id, inv_id):
    try:
        inv_res = db.table("invullingen").select("*").eq("id", inv_id).eq("dossier_id", dossier_id).single().execute()
    except Exception:
        abort(404)

    if not inv_res.data:
        abort(404)

    inv = inv_res.data
    waarden = inv.get("waarden") or {}
    if isinstance(waarden, str):
        waarden = json.loads(waarden)

    try:
        tmpl_res = db.table("templates").select("*").eq("id", inv["template_id"]).single().execute()
    except Exception:
        abort(404)

    template = tmpl_res.data
    fields = template.get("fields") or []
    if isinstance(fields, str):
        fields = json.loads(fields)
    waarden = format_field_values(fields, waarden)

    # Haal dossier + alle invullingen op voor systeemvariabelen
    try:
        dos_res = db.table("dossiers").select("*").eq("id", dossier_id).single().execute()
        alle_inv = db.table("invullingen").select("template_id").eq("dossier_id", dossier_id).execute()
        template_ids = [i["template_id"] for i in (alle_inv.data or [])]
        alle_tmpl = db.table("templates").select("id,name").in_("id", template_ids).execute()
        tmpl_map = {t["id"]: t for t in (alle_tmpl.data or [])}
        templates_in_dossier = [tmpl_map[tid] for tid in template_ids if tid in tmpl_map]
        positie = next((i + 1 for i, tid in enumerate(template_ids) if tid == inv["template_id"]), None)
        dossier = dos_res.data or {}
    except Exception:
        templates_in_dossier = [template]
        positie = 1
        dossier = {}

    cfg = get_config()
    system_vals = build_system_values(
        cfg,
        dossier=dossier,
        templates_in_dossier=templates_in_dossier,
        positie=positie,
        user_label=session.get("user_name", session.get("user_email", "")),
    )
    gedeelde = dossier.get("gedeelde_waarden") or {}
    if isinstance(gedeelde, str):
        gedeelde = json.loads(gedeelde)
    # Volgorde: systeem < gedeeld (dossier-niveau) < per-invulling
    waarden = {**system_vals, **gedeelde, **waarden}

    try:
        docx_bytes = supabase.storage.from_(SUPABASE_BUCKET).download(template["docx_path"])
    except Exception as e:
        flash(f"Fout bij ophalen sjabloonbestand: {e}", "error")
        return redirect(url_for("dossier_detail", dossier_id=dossier_id))

    try:
        result = fill_template(docx_bytes, waarden)
    except Exception as e:
        flash(f"Fout bij genereren document: {e}", "error")
        return redirect(url_for("dossier_detail", dossier_id=dossier_id))

    filename = f"{template['name'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return send_file(
        io.BytesIO(result),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename,
    )


@app.route("/dossier/<dossier_id>/invulling/<inv_id>/toegang", methods=["POST"])
@login_required
def dossier_invulling_toegang(dossier_id, inv_id):
    extern_toegang = request.form.get("extern_toegang", "verborgen")
    if extern_toegang not in ("verborgen", "leesbaar", "invulbaar"):
        extern_toegang = "verborgen"

    try:
        db.table("invullingen").update({
            "extern_toegang": extern_toegang,
            "updated_at": datetime.utcnow().isoformat(),
        }).eq("id", inv_id).eq("dossier_id", dossier_id).execute()
        flash("Toegang bijgewerkt.", "success")
    except Exception as e:
        flash(f"Fout bij opslaan toegang: {e}", "error")

    return redirect(url_for("dossier_detail", dossier_id=dossier_id))


@app.route("/dossier/<dossier_id>/status", methods=["POST"])
@login_required
def dossier_status(dossier_id):
    status = request.form.get("status", "concept")
    if status not in ("concept", "afgerond"):
        status = "concept"

    try:
        db.table("dossiers").update({
            "status": status,
            "updated_at": datetime.utcnow().isoformat(),
        }).eq("id", dossier_id).execute()
        flash("Dossier status bijgewerkt.", "success")
    except Exception as e:
        flash(f"Fout bij bijwerken status: {e}", "error")

    return redirect(url_for("dossier_detail", dossier_id=dossier_id))


@app.route("/dossier/<dossier_id>/bewerken", methods=["POST"])
@login_required
def dossier_bewerken(dossier_id):
    naam = request.form.get("naam", "").strip()
    omschrijving = request.form.get("omschrijving", "").strip()
    jaar_raw = request.form.get("jaar", "").strip()
    vormen = request.form.getlist("financieringsvorm")
    financieringsvorm = ", ".join(sorted(v.strip() for v in vormen if v.strip())) or None
    oa_type = request.form.get("oa_type", "").strip() or None

    try:
        jaar = int(jaar_raw) if jaar_raw else None
    except ValueError:
        jaar = None

    if not naam:
        flash("Naam is verplicht.", "error")
        return redirect(url_for("dossier_detail", dossier_id=dossier_id))

    try:
        db.table("dossiers").update({
            "naam": naam,
            "omschrijving": omschrijving or None,
            "jaar": jaar,
            "financieringsvorm": financieringsvorm,
            "oa_type": oa_type,
            "updated_at": datetime.utcnow().isoformat(),
        }).eq("id", dossier_id).execute()
        flash("Dossier bijgewerkt.", "success")
    except Exception as e:
        flash(f"Fout bij opslaan: {e}", "error")

    return redirect(url_for("dossier_detail", dossier_id=dossier_id))


@app.route("/dossier/<dossier_id>/verwijderen", methods=["POST"])
@login_required
def dossier_verwijderen(dossier_id):
    try:
        # Verwijder gekoppelde tokens en invullingen eerst (foreign key)
        db.table("dossier_tokens").delete().eq("dossier_id", dossier_id).execute()
        db.table("invullingen").delete().eq("dossier_id", dossier_id).execute()
        db.table("dossiers").delete().eq("id", dossier_id).execute()
        flash("Dossier verwijderd.", "success")
    except Exception as e:
        flash(f"Fout bij verwijderen: {e}", "error")
        return redirect(url_for("dossier_detail", dossier_id=dossier_id))
    return redirect(url_for("dossiers_overzicht"))


@app.route("/dossier/<dossier_id>/invulling/<inv_id>/heropenen", methods=["POST"])
@login_required
def dossier_invulling_heropenen(dossier_id, inv_id):
    try:
        db.table("invullingen").update({
            "extern_status": "open",
            "updated_at": datetime.utcnow().isoformat(),
        }).eq("id", inv_id).eq("dossier_id", dossier_id).execute()
        flash("Invulling heropend — de externe partij kan opnieuw aanpassen.", "success")
    except Exception as e:
        flash(f"Fout bij heropenen: {e}", "error")
    return redirect(url_for("dossier_detail", dossier_id=dossier_id))


# ---------------------------------------------------------------------------
# User preferences API
# ---------------------------------------------------------------------------

@app.route("/api/preferences", methods=["GET"])
@login_required
def get_preferences():
    user_id = session.get("user_id")
    try:
        res = db.table("user_preferences").select("preferences").eq("user_id", user_id).execute()
        prefs = res.data[0]["preferences"] if res.data else {}
    except Exception:
        prefs = {}
    return {"preferences": prefs}


@app.route("/api/preferences", methods=["POST"])
@login_required
def set_preferences():
    user_id = session.get("user_id")
    data = request.get_json(silent=True) or {}
    nieuwe_prefs = data.get("preferences", {})
    if not isinstance(nieuwe_prefs, dict):
        return {"error": "Ongeldig formaat"}, 400
    try:
        # Haal bestaande preferences op en merge
        res = db.table("user_preferences").select("preferences").eq("user_id", user_id).execute()
        bestaand = res.data[0]["preferences"] if res.data else {}
        if not isinstance(bestaand, dict):
            bestaand = {}
        samengevoegd = {**bestaand, **nieuwe_prefs}
        db.table("user_preferences").upsert(
            {"user_id": user_id, "preferences": samengevoegd},
            on_conflict="user_id",
        ).execute()
    except Exception as e:
        app.logger.error(f"Fout bij opslaan preferences user={user_id}: {e}")
        return {"error": str(e)}, 500
    return {"ok": True}


@app.route("/dossier/<dossier_id>/token/aanmaken", methods=["POST"])
@login_required
def dossier_token_aanmaken(dossier_id):
    omschrijving = request.form.get("omschrijving", "").strip() or "Extern dossier"
    try:
        db.table("dossier_tokens").insert({
            "dossier_id": dossier_id,
            "omschrijving": omschrijving,
            "status": "actief",
        }).execute()
        flash("Deellink aangemaakt.", "success")
    except Exception as e:
        flash(f"Fout bij aanmaken deellink: {e}", "error")

    return redirect(url_for("dossier_detail", dossier_id=dossier_id))


@app.route("/dossier_token/<token_id>/intrekken", methods=["POST"])
@login_required
def dossier_token_intrekken(token_id):
    try:
        tok = db.table("dossier_tokens").select("dossier_id").eq("id", token_id).single().execute()
        db.table("dossier_tokens").update({"status": "ingetrokken"}).eq("id", token_id).execute()
        flash("Deellink ingetrokken.", "success")
        if tok.data:
            return redirect(url_for("dossier_detail", dossier_id=tok.data["dossier_id"]))
    except Exception as e:
        flash(f"Fout bij intrekken: {e}", "error")
    return redirect(url_for("dossiers_overzicht"))


@app.route("/dossier_token/<token_id>/verwijderen", methods=["POST"])
@login_required
def dossier_token_verwijderen(token_id):
    try:
        tok = db.table("dossier_tokens").select("dossier_id").eq("id", token_id).single().execute()
        dossier_id = tok.data["dossier_id"] if tok.data else None
        db.table("dossier_tokens").delete().eq("id", token_id).execute()
        flash("Deellink verwijderd.", "success")
        if dossier_id:
            return redirect(url_for("dossier_detail", dossier_id=dossier_id))
    except Exception as e:
        flash(f"Fout bij verwijderen: {e}", "error")
    return redirect(url_for("dossiers_overzicht"))


# ---------------------------------------------------------------------------
# Extern dossier routes (geen login)
# ---------------------------------------------------------------------------

@app.route("/dossier/extern/<token_id>", methods=["GET", "POST"])
def dossier_extern(token_id):
    cfg = get_config()
    try:
        tok_res = db.table("dossier_tokens").select("*").eq("id", token_id).single().execute()
    except Exception:
        abort(404)

    if not tok_res.data:
        abort(404)

    token = tok_res.data
    if token["status"] != "actief":
        abort(404)

    dossier_id = token["dossier_id"]
    try:
        dos_res = db.table("dossiers").select("*").eq("id", dossier_id).single().execute()
    except Exception:
        abort(404)

    if not dos_res.data:
        abort(404)

    dossier = dos_res.data

    # Dossier afgesloten: externe pagina altijd vergrendeld
    dossier_afgesloten = dossier.get("status") == "afgerond"

    try:
        inv_res = db.table("invullingen").select("*").eq("dossier_id", dossier_id).execute()
        invullingen_raw = inv_res.data or []
    except Exception:
        invullingen_raw = []

    # Enrich with template fields; skip verborgen
    invullingen = []
    for inv in invullingen_raw:
        if inv.get("extern_toegang", "verborgen") == "verborgen":
            continue
        try:
            tmpl_res = db.table("templates").select("*").eq("id", inv["template_id"]).single().execute()
            tmpl = tmpl_res.data or {}
        except Exception:
            tmpl = {}
        all_fields = tmpl.get("fields") or []
        if isinstance(all_fields, str):
            all_fields = json.loads(all_fields)
        waarden = inv.get("waarden") or {}
        if isinstance(waarden, str):
            waarden = json.loads(waarden)
        # Externe pagina toont altijd alleen velden met eigenaar == "extern"
        visible_fields = [f for f in all_fields if f.get("eigenaar") == "extern"]
        # Invulling vergrendeld als dossier afgerond is of extern_status == verzonden
        invulling_vergrendeld = (
            dossier_afgesloten or
            inv.get("extern_status") == "verzonden"
        )
        # Overschrijf toegang naar leesbaar als vergrendeld maar niet verborgen
        effective_toegang = inv.get("extern_toegang", "verborgen")
        if invulling_vergrendeld and effective_toegang == "invulbaar":
            effective_toegang = "leesbaar"

        invullingen.append({
            **inv,
            "template": tmpl,
            "fields": visible_fields,
            "waarden": waarden,
            "effective_toegang": effective_toegang,
            "vergrendeld": invulling_vergrendeld,
        })

    if request.method == "POST":
        # Blokkeer POST als dossier afgerond
        if dossier_afgesloten:
            flash("Dit dossier is afgesloten en kan niet meer worden bewerkt.", "error")
            return redirect(url_for("dossier_extern", token_id=token_id))

        errors = []
        for inv in invullingen:
            if inv["extern_toegang"] != "invulbaar":
                continue
            if inv["vergrendeld"]:
                continue
            existing_waarden = inv.get("waarden") or {}
            new_waarden = dict(existing_waarden)
            for field in inv["fields"]:
                key = field["name"]
                val = request.form.get(f"{inv['id']}_{key}", "").strip()
                if field.get("required") and not val:
                    errors.append(f'"{field.get("label", key)}" is verplicht.')
                new_waarden[key] = val

            if not errors:
                try:
                    db.table("invullingen").update({
                        "waarden": new_waarden,
                        "extern_status": "verzonden",
                        "updated_at": datetime.utcnow().isoformat(),
                    }).eq("id", inv["id"]).execute()
                except Exception as e:
                    errors.append(f"Fout bij opslaan: {e}")

        if errors:
            for err in errors:
                flash(err, "error")
            return render_template(
                "dossier_extern.html",
                cfg=cfg,
                dossier=dossier,
                dossier_afgesloten=dossier_afgesloten,
                invullingen=invullingen,
                token=token,
                prefill=request.form,
            )

        return redirect(url_for("dossier_extern_bedankt", token_id=token_id))

    return render_template(
        "dossier_extern.html",
        cfg=cfg,
        dossier=dossier,
        dossier_afgesloten=dossier_afgesloten,
        invullingen=invullingen,
        token=token,
        prefill={},
    )


@app.route("/dossier/extern/<token_id>/bedankt")
def dossier_extern_bedankt(token_id):
    cfg = get_config()
    try:
        tok_res = db.table("dossier_tokens").select("dossier_id").eq("id", token_id).single().execute()
        if tok_res.data:
            dos_res = db.table("dossiers").select("naam").eq("id", tok_res.data["dossier_id"]).single().execute()
            dossier_naam = dos_res.data["naam"] if dos_res.data else ""
        else:
            dossier_naam = ""
    except Exception:
        dossier_naam = ""
    return render_template("dossier_extern_bedankt.html", cfg=cfg, dossier_naam=dossier_naam)


# ---------------------------------------------------------------------------
# Inzendingen (voor alle ingelogde gebruikers)
# ---------------------------------------------------------------------------

@app.route("/inzendingen")
@login_required
def inzendingen_overzicht():
    cfg = get_config()

    # Haal alle invullingen op waarbij de externe partij iets heeft ingevuld
    # (extern_toegang != verborgen én waarden niet leeg)
    try:
        inv_res = db.table("invullingen").select("*").neq("extern_toegang", "verborgen").execute()
        alle_invullingen = inv_res.data or []
    except Exception as e:
        flash(f"Fout bij ophalen inzendingen: {e}", "error")
        alle_invullingen = []

    # Filter op invullingen met minstens één extern ingevuld veld
    inzendingen = []
    dossier_cache = {}
    template_cache = {}

    for inv in alle_invullingen:
        waarden = inv.get("waarden") or {}
        if isinstance(waarden, str):
            waarden = json.loads(waarden)

        if not any(str(v).strip() for v in waarden.values()):
            continue

        dos_id = inv.get("dossier_id")
        if dos_id not in dossier_cache:
            try:
                dos_res = db.table("dossiers").select("naam,status").eq("id", dos_id).single().execute()
                dossier_cache[dos_id] = dos_res.data or {}
            except Exception:
                dossier_cache[dos_id] = {}

        tmpl_id = inv.get("template_id")
        if tmpl_id not in template_cache:
            try:
                tmpl_res = db.table("templates").select("name,fields").eq("id", tmpl_id).single().execute()
                template_cache[tmpl_id] = tmpl_res.data or {}
            except Exception:
                template_cache[tmpl_id] = {}

        tmpl = template_cache.get(tmpl_id, {})
        fields = tmpl.get("fields") or []
        if isinstance(fields, str):
            fields = json.loads(fields)

        # Alleen extern-ingevulde velden tonen
        extern_velden = [f for f in fields if f.get("eigenaar") == "extern"]

        inzendingen.append({
            **inv,
            "waarden": waarden,
            "dossier": dossier_cache.get(dos_id, {}),
            "template_naam": tmpl.get("name", "—"),
            "extern_velden": extern_velden,
        })

    # Sorteer: meest recent bijgewerkt bovenaan
    inzendingen.sort(key=lambda x: x.get("updated_at") or x.get("created_at") or "", reverse=True)

    # Markeer als gelezen voor deze gebruiker
    user_id = session.get("user_id")
    if user_id:
        try:
            db.table("inzendingen_gelezen").upsert({
                "user_id": user_id,
                "gezien_op": datetime.utcnow().isoformat(),
            }, on_conflict="user_id").execute()
        except Exception:
            pass

    return render_template("inzendingen.html", cfg=cfg, inzendingen=inzendingen)


# ---------------------------------------------------------------------------
# Error handlers
# ---------------------------------------------------------------------------

@app.errorhandler(404)
def not_found(e):
    cfg = get_config()
    return render_template("base.html", cfg=cfg, error="Pagina niet gevonden (404)"), 404


@app.errorhandler(500)
def server_error(e):
    cfg = get_config()
    return render_template("base.html", cfg=cfg, error="Interne serverfout (500)"), 500


if __name__ == "__main__":
    app.run(debug=True)

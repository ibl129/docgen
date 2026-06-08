"""Regressietests voor de docx-invulengine in app.py.

Bewaakt de vier if-/placeholder-fixes van 8 jun 2026 (zie geheugen-notitie
docgen_docx_engine_valkuilen). Draait zonder Flask: de relevante functies worden
uit app.py geïsoleerd en los uitgevoerd, zodat zware web-deps (whitenoise e.d.)
niet nodig zijn.

Run:  python3 tests/test_docx_render.py
(exit-code 0 = alle tests groen)
"""

import ast
import io
import os
import re as _re

from docx import Document

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
FIX = os.path.join(HERE, "fixtures")


def _load_docx_functions():
    """Voer alleen de docx-gerelateerde functies uit app.py uit, zonder Flask te importeren."""
    src = open(os.path.join(ROOT, "app.py"), encoding="utf-8").read()
    mod = ast.parse(src)
    wanted = {
        "_merge_placeholder_runs", "_replace_in_paragraph", "_process_textarea",
        "_process_conditionals", "_strip_tag_from_para", "_para_set_text",
        "_para_full_text", "_remove_para_spacing", "_insert_paragraph_after",
        "fill_template",
    }
    glb = {}
    exec(
        "from docx.oxml.ns import qn\n"
        "import re as _re\nimport re\nimport io\n"
        "from docx import Document\n"
        "from docx.oxml import OxmlElement\n"
        "from docx.text.paragraph import Paragraph as DocxParagraph\n"
        "from copy import deepcopy\n"
        "import lxml.etree as _etree\n"
        "from datetime import datetime\n",
        glb,
    )
    # Module-level regexen (_IF_RE etc.) meenemen.
    for name in ("_ENDIF_RE", "_IF_RE", "_IFNOT_RE"):
        m = _re.search(rf"{name}\s*=\s*_re\.compile\([^\n]+\)", src)
        if m:
            exec(m.group(0), glb)
    for node in mod.body:
        if isinstance(node, ast.FunctionDef) and node.name in wanted:
            exec(ast.get_source_segment(src, node), glb)
    return glb


_FN = _load_docx_functions()


def render(template_name, values):
    with open(os.path.join(FIX, template_name), "rb") as f:
        data = f.read()
    return Document(io.BytesIO(_FN["fill_template"](data, values)))


def texts(doc):
    return [p.text for p in doc.paragraphs if p.text.strip()]


def has(doc, fragment):
    return any(fragment in p.text for p in doc.paragraphs)


def element_seq(doc, fragment):
    """Geef de volgorde van [TAB]/tekst-elementen van de eerste alinea die `fragment` bevat."""
    for p in doc.paragraphs:
        if fragment in p.text:
            seq = []
            for el in p._element.iter():
                tag = el.tag.split("}")[-1]
                if tag == "tab":
                    seq.append("[TAB]")
                elif tag == "t" and (el.text or ""):
                    seq.append(el.text)
            return seq
    return None


# --- Gedeelde basiswaarden -------------------------------------------------

BIJLAGE1 = {
    "datum_aanvang": "1-1-2026", "datum_einde": "31-12-2026",
    "zvw_max_uren_per_nacht": "8", "uurtarief_start_einddag": "67.22",
    "wlz_max_uren_per_nacht": "8",
    "tarief_persoonlijke_verzorging": "63.99", "tarief_verpleging": "70.50",
    "contractjaar": "2026", "contractnummer": "123",
    "aantal_uren_per_kwartaal": "1500",
    "opdrachtgever_naam_ondertekening": "A", "opdrachtnemer_naam_ondertekening": "B",
    "opdrachtgever_functie": "f1", "opdrachtnemer_functie": "f2",
    "opdrachtgever_naam": "n1", "opdrachtnemer_naam": "n2",
}

OVEREENKOMST = {
    "_fin_zvw": "ja", "_fin_wlz": "ja", "_fin_wmo": "",
    "specifieke_kennis": "wijkverpleging", "invulling_specifieke_zorg": "",
    "exclusieve_samenwerking": "", "wlz_zorgkantoren": "X", "wmo_gemeenten": "Y",
    "lijst_zorgverzekeraars": "Z", "wlz_zorgsoort": "s", "wlz_regio": "r",
    "afstemmen_aantal_wlz_uren": "",
}


# --- Tests -----------------------------------------------------------------

def test_dubbele_placeholder_per_alinea():
    """FIX 1: dezelfde {{placeholder}} kan meermaals in één alinea voorkomen."""
    doc = render("bijlage1_tarief.docx", BIJLAGE1)
    zin = next((p.text for p in doc.paragraphs if "realistische afspraak" in p.text), "")
    assert "{{contractjaar}}" not in zin, "tweede {{contractjaar}} in dezelfde alinea niet vervangen"
    assert zin.rstrip().endswith("2026."), f"verwachtte '...2026.', kreeg: {zin[-40:]!r}"


def test_tab_blijft_op_plek_in_if():
    """FIX 2: een <w:tab/> tussen tekst en {{placeholder}} binnen een {{#if}} blijft staan."""
    doc = render("bijlage1_tarief.docx", BIJLAGE1)
    seq = element_seq(doc, "Persoonlijke verzorging")
    assert seq is not None
    # Verwacht: ... 'Persoonlijke verzorging:'  [TAB]  '€'  '63.99'
    label_idx = next(i for i, x in enumerate(seq) if "Persoonlijke verzorging" in x)
    rest = seq[label_idx:]
    assert "[TAB]" in rest, f"tab tussen label en bedrag verdwenen: {rest!r}"
    assert any("63.99" in x for x in rest), "tarief niet ingevuld"
    # De tab moet vóór het euroteken/bedrag staan, niet erna.
    tab_pos = rest.index("[TAB]")
    euro_pos = next(i for i, x in enumerate(rest) if "€" in x or "63.99" in x)
    assert tab_pos < euro_pos, f"tab staat op verkeerde plek: {rest!r}"


def test_geen_runaway_delete_bij_foute_endif():
    """FIX 3: typefout {{//if}} mag niet het halve document wissen."""
    doc = render("bijlage1_tarief.docx", BIJLAGE1)
    # Het Bijlage-1-template bevat een {{//if}} (dubbele slash) op de Verpleging-regel.
    assert has(doc, "Verpleging"), "Verpleging-regel (na {{//if}}) ontbreekt — runaway delete?"
    assert has(doc, "Datum:"), "einde van het document ontbreekt — runaway delete?"
    assert has(doc, "70.50"), "tarief_verpleging niet ingevuld"
    # Geen restanten van tags.
    for p in doc.paragraphs:
        assert "{{#if" not in p.text and "{{/if" not in p.text and "{{//if" not in p.text, \
            f"resterende if-tag: {p.text[:60]!r}"


def test_multi_alinea_if_behoudt_tekst_naast_tag():
    """FIX 4: tekst naast de openingstag van een multi-alinea {{#if}} blijft bij keep behouden."""
    doc = render("overeenkomst.docx", OVEREENKOMST)
    assert has(doc, "a. Cliënt met wijkverpleging"), \
        "kop 'a. Cliënt met wijkverpleging:' verdwenen (stond naast {{#if _fin_zvw}})"
    assert has(doc, "b. Cliënt met indicatie"), \
        "kop 'b. Cliënt met indicatie...' verdwenen (stond naast {{#if _fin_wlz}})"
    for p in doc.paragraphs:
        assert "{{#if" not in p.text and "{{/if" not in p.text, f"resterende if-tag: {p.text[:60]!r}"


def test_multi_alinea_if_verwijdert_bij_lege_waarde():
    """Tegenproef bij FIX 4: leeg veld → blok (incl. kop) verdwijnt, rest blijft intact."""
    vals = dict(OVEREENKOMST, _fin_zvw="")
    doc = render("overeenkomst.docx", vals)
    assert not has(doc, "a. Cliënt met wijkverpleging"), "blok 'a.' had weg moeten zijn (_fin_zvw leeg)"
    assert has(doc, "b. Cliënt met indicatie"), "blok 'b.' had moeten blijven (_fin_wlz gevuld)"
    assert not has(doc, "ingevolge de Zorgverzekeringswet"), "zvw-considerans had weg moeten zijn"
    assert has(doc, "Vergoeding en betaling"), "document niet intact tot het einde"


def test_eenregelig_if_persoonlijke_verzorging_leeg():
    """Eenregelig {{#if}}-blok (patroon 1): leeg veld verbergt alleen die regel."""
    vals = dict(BIJLAGE1, tarief_persoonlijke_verzorging="")
    doc = render("bijlage1_tarief.docx", vals)
    assert not any(t.strip().startswith("Persoonlijke verzorging") for t in texts(doc)), \
        "lege persoonlijke verzorging niet verborgen"
    assert has(doc, "Verpleging"), "andere regel ten onrechte verborgen"


def _main():
    tests = [v for k, v in sorted(globals().items()) if k.startswith("test_") and callable(v)]
    failed = 0
    for t in tests:
        try:
            t()
            print(f"  PASS  {t.__name__}")
        except AssertionError as e:
            failed += 1
            print(f"  FAIL  {t.__name__}: {e}")
        except Exception as e:  # noqa: BLE001
            failed += 1
            print(f"  ERROR {t.__name__}: {type(e).__name__}: {e}")
    print(f"\n{len(tests) - failed}/{len(tests)} geslaagd")
    return 1 if failed else 0


if __name__ == "__main__":
    raise SystemExit(_main())
